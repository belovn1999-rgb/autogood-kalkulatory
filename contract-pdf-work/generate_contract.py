#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import shutil
import tempfile
import zipfile
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from lxml import etree


ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "templates" / "Umowa_Zamowienia_Pojazdu_AG_template.docx"
SIGNATURE_STAMP = ROOT / "assets" / "signature_stamp_source.jpg"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
NS = {"w": W_NS, "w14": W14_NS}

POLISH_MONTHS = {
    1: "stycznia",
    2: "lutego",
    3: "marca",
    4: "kwietnia",
    5: "maja",
    6: "czerwca",
    7: "lipca",
    8: "sierpnia",
    9: "września",
    10: "października",
    11: "listopada",
    12: "grudnia",
}


CHECKBOX_INDEX = {
    "client_is_entrepreneur": 1,
    "subject_mediation": 2,
    "subject_purchase_by_autogood": 3,
    "subject_financing": 4,
    "subject_client_indicated_vehicle": 5,
    "fuel_diesel": 6,
    "fuel_benzyna": 7,
    "fuel_hybryda": 8,
    "fuel_elektryk": 9,
    "gearbox_manualna": 10,
    "gearbox_automatyczna": 11,
    "euro_6": 12,
    "euro_7": 13,
    "euro_dowolna": 14,
    "euro_inna": 15,
    "body_sedan": 16,
    "body_kombi": 17,
    "body_coupe": 18,
    "body_inne": 19,
    "allow_collision": 20,
}


def contract_number(contract_date: date, sequence: int = 1) -> str:
    base = f"{contract_date.day}/{contract_date.month}/{str(contract_date.year)[-2:]}"
    if sequence and sequence > 1:
        return f"{base}/{sequence}"
    return base


def polish_date_line(contract_date: date) -> str:
    month = POLISH_MONTHS[contract_date.month]
    return f"Łomianki, {contract_date.day} {month} {contract_date.year} roku"


def parse_date(value: str) -> date:
    return datetime.strptime(value, "%Y-%m-%d").date()


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def set_paragraph_text(paragraph, text: str, *, size_pt: float = 8, bold: bool = False) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.bold = bold


def append_text(paragraph, text: str, *, size_pt: float = 8, bold: bool = False, line_break: bool = False) -> None:
    run = paragraph.add_run()
    if line_break:
        run.add_break()
    run.add_text(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_empty_paragraph(paragraph, text: str, *, size_pt: float = 8, bold: bool = False) -> None:
    clear_paragraph(paragraph)
    append_text(paragraph, text, size_pt=size_pt, bold=bold)


def checked_keys(data: dict) -> set[str]:
    checks: set[str] = set()
    client = data["client"]
    if client.get("is_entrepreneur", client["type"] == "company"):
        checks.add("client_is_entrepreneur")

    subject = data["agreement"].get("subject")
    if subject == "mediation":
        checks.add("subject_mediation")
    elif subject == "purchase_by_autogood":
        checks.add("subject_purchase_by_autogood")
    elif subject == "financing":
        checks.add("subject_financing")

    if data["agreement"].get("client_indicated_vehicle"):
        checks.add("subject_client_indicated_vehicle")

    vehicle = data["vehicle"]
    for fuel in vehicle.get("fuel") or []:
        checks.add(f"fuel_{fuel}")

    gearbox = vehicle.get("gearbox")
    if gearbox:
        checks.add(f"gearbox_{gearbox}")

    euro = vehicle.get("euro_standard")
    if euro:
        checks.add(f"euro_{euro}")

    body_type = (vehicle.get("body") or {}).get("type")
    if body_type:
        checks.add(f"body_{body_type}")

    if vehicle.get("allow_collision_without_longitudinals"):
        checks.add("allow_collision")

    return checks


def set_docx_checkboxes(docx_path: Path, checked: set[str]) -> None:
    checked_numbers = {CHECKBOX_INDEX[key] for key in checked if key in CHECKBOX_INDEX}

    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td)
        with zipfile.ZipFile(docx_path) as zf:
            zf.extractall(tmp)

        document_xml = tmp / "word" / "document.xml"
        tree = etree.parse(str(document_xml))
        root = tree.getroot()
        checkboxes = root.xpath(".//w:sdt[w:sdtPr/w14:checkbox]", namespaces=NS)

        for idx, sdt in enumerate(checkboxes, start=1):
            is_checked = idx in checked_numbers
            checked_el = sdt.xpath("./w:sdtPr/w14:checkbox/w14:checked", namespaces=NS)
            if checked_el:
                checked_el[0].set(f"{{{W14_NS}}}val", "1" if is_checked else "0")

            text_el = sdt.xpath("./w:sdtContent//w:t", namespaces=NS)
            if text_el:
                text_el[0].text = "þ" if is_checked else "¨"

        tree.write(str(document_xml), xml_declaration=True, encoding="UTF-8", standalone="yes")

        backup = docx_path.with_suffix(".docx.bak")
        shutil.copy2(docx_path, backup)
        with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for file in tmp.rglob("*"):
                if file.is_file():
                    zf.write(file, file.relative_to(tmp))
        backup.unlink(missing_ok=True)


def fill_docx(data: dict, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE))

    contract_date = parse_date(data["contract"]["date"])
    number = contract_number(contract_date, int(data["contract"].get("sequence", 1)))

    set_paragraph_text(doc.paragraphs[0], polish_date_line(contract_date), size_pt=14)
    doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_text(doc.paragraphs[1], f"UMOWA ZAMÓWIENIA POJAZDU {number}", size_pt=17, bold=True)

    table = doc.tables[0]
    client = data["client"]
    budget = data["budget"]
    vehicle = data["vehicle"]

    append_text(table.cell(2, 0).paragraphs[1], client["name"], size_pt=8, bold=True, line_break=True)
    set_empty_paragraph(table.cell(3, 0).paragraphs[2], client["address"], size_pt=8)

    if client["type"] == "company":
        identifier = client.get("nip", "")
        document_text = ""
    else:
        identifier = client.get("pesel", "")
        document_text = client.get("document", "")

    set_empty_paragraph(table.cell(4, 0).paragraphs[2], identifier, size_pt=8)
    set_empty_paragraph(table.cell(5, 0).paragraphs[2], document_text, size_pt=8)
    append_text(table.cell(6, 0).paragraphs[1], client["phone"], size_pt=8, line_break=False)
    append_text(table.cell(6, 0).paragraphs[2], client["email"], size_pt=8, line_break=False)

    set_empty_paragraph(table.cell(7, 3).paragraphs[2], vehicle["make_model"], size_pt=8, bold=True)
    append_text(table.cell(10, 3).paragraphs[2], vehicle["first_registration"], size_pt=8)
    append_text(table.cell(11, 3).paragraphs[2], f" {vehicle['mileage_to']}", size_pt=8)

    append_text(table.cell(11, 0).paragraphs[1], budget["total"], size_pt=8, bold=True)
    append_text(table.cell(12, 0).paragraphs[1], budget["advance"], size_pt=8, bold=True)

    body = vehicle.get("body") or {}
    if body.get("type") == "inne" and body.get("other"):
        append_text(table.cell(12, 3).paragraphs[0], body["other"], size_pt=8)

    set_empty_paragraph(table.cell(13, 3).paragraphs[2], vehicle.get("required_equipment", ""), size_pt=8)
    set_empty_paragraph(table.cell(15, 3).paragraphs[2], vehicle.get("expected_equipment", ""), size_pt=8)

    if SIGNATURE_STAMP.exists():
        sig_paragraph = table.cell(18, 3).paragraphs[1]
        sig_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig_paragraph.add_run().add_picture(str(SIGNATURE_STAMP), width=Inches(1.55))

    doc.save(str(output))
    set_docx_checkboxes(output, checked_keys(data))


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate AUTOGOOD contract DOCX from JSON data.")
    parser.add_argument("data", type=Path)
    parser.add_argument("--output", type=Path, default=ROOT / "output" / "contract.docx")
    args = parser.parse_args()

    data = json.loads(args.data.read_text(encoding="utf-8"))
    fill_docx(data, args.output)
    print(args.output)


if __name__ == "__main__":
    main()
